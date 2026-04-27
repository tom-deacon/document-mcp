"""Document parsing — plain-text fallback + optional structured path."""

import hashlib
import logging
import mimetypes
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

import fitz  # PyMuPDF
import tiktoken
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parse documents of various formats.

    When ``use_structured=True`` (the default), DOCX / PDF / Excel files are
    parsed through the new structured pipeline that preserves headings, tables,
    and sheet layout.  Plain-text files (.txt / .md / .rtf) always use the
    original character-window chunker — they have no structure to preserve.

    Setting ``use_structured=False`` or calling any of the private
    ``_extract_*`` methods reverts to the original flat-text behaviour, so
    existing callers are fully backward-compatible.
    """

    def __init__(
        self,
        chunk_size: int = 2000,
        chunk_overlap: int = 200,
        use_structured: bool = True,
        max_chunk_tokens: int = 512,
        enable_ocr: bool = True,
        ocr_language: str = "eng",
        enable_vision_summary: bool = False,
        vision_model: str = "",
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.use_structured = use_structured
        self.max_chunk_tokens = max_chunk_tokens
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
        # Build a VisualConfig so parsers can share it without re-importing config
        from .parsers.visual_utils import VisualConfig
        self._visual_config = VisualConfig(
            enable_ocr=enable_ocr,
            ocr_language=ocr_language,
            enable_vision_summary=enable_vision_summary,
            vision_model=vision_model,
        )

    # ------------------------------------------------------------------
    # Public entry point
    # ------------------------------------------------------------------

    def parse_file(self, file_path: Path) -> Dict[str, Any]:
        """Parse a file and return metadata, full text, and chunks.

        Return shape (unchanged from original):
          metadata, text, chunks[], num_chunks, total_chars, total_tokens

        Each chunk now carries additional optional keys when the structured
        path is used:
          element_type, section_path, page_or_sheet, embedding_text,
          html_or_markdown
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        file_stats = file_path.stat()
        file_hash = self._calculate_file_hash(file_path)
        file_type = file_path.suffix.lower()

        metadata = {
            "file_path": str(file_path),
            "file_name": file_path.name,
            "file_size": file_stats.st_size,
            "modified_time": datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
            "file_hash": file_hash,
            "file_type": file_type,
        }

        # Route to structured or plain-text pipeline
        if self.use_structured and file_type in _STRUCTURED_TYPES:
            logger.info(
                "    -> [STRUCTURED] Parsing %s as %s", file_path.name, file_type
            )
            chunks, full_text = self._parse_structured(file_path, file_type)
        else:
            logger.info(
                "    -> [PLAIN-TEXT] Parsing %s as %s", file_path.name, file_type
            )
            full_text = self._extract_text(file_path)
            logger.info("    -> Extracted %s characters", f"{len(full_text):,}")
            chunks = self._create_chunks(full_text)

        total_tokens = len(self.tokenizer.encode(full_text)) if full_text else 0

        return {
            "metadata": metadata,
            "text": full_text,
            "chunks": chunks,
            "num_chunks": len(chunks),
            "total_chars": len(full_text),
            "total_tokens": total_tokens,
        }

    # ------------------------------------------------------------------
    # Structured pipeline
    # ------------------------------------------------------------------

    def _parse_structured(
        self, file_path: Path, file_type: str
    ) -> tuple[List[Dict[str, Any]], str]:
        """Run the structured parser + chunker for the given file type.

        Returns (chunks, full_text).  Falls back to plain-text on any error.
        """
        try:
            elements = _route_to_parser(file_path, file_type, self._visual_config)
        except Exception as exc:
            logger.warning(
                "    -> Structured parser failed for %s (%s). "
                "Falling back to plain-text extraction.",
                file_path.name, exc,
            )
            full_text = self._extract_text(file_path)
            return self._create_chunks(full_text), full_text

        if not elements:
            logger.warning(
                "    -> Structured parser returned no elements for %s. "
                "Falling back to plain-text extraction.",
                file_path.name,
            )
            full_text = self._extract_text(file_path)
            return self._create_chunks(full_text), full_text

        # Reconstruct a plain-text version from elements (for summary / LLM)
        full_text = "\n\n".join(e.get("text", "") for e in elements if e.get("text"))

        # Chunk the elements with the structure-aware chunker
        from .chunkers.structure_aware import chunk_structured_elements

        document_title = file_path.stem  # filename without extension
        chunks = chunk_structured_elements(
            elements,
            document_title=document_title,
            max_tokens=self.max_chunk_tokens,
            tokenizer=self.tokenizer,
        )

        logger.info(
            "    -> Structured parse: %d elements -> %d chunks",
            len(elements), len(chunks),
        )

        # Log quality warning if very few chunks came out of a non-empty file
        if len(chunks) == 0 and len(full_text) > 200:
            logger.warning(
                "    -> WARNING: 0 chunks produced for non-empty file %s — "
                "check parser output",
                file_path.name,
            )

        return chunks, full_text

    # ------------------------------------------------------------------
    # Plain-text extraction (unchanged from original)
    # ------------------------------------------------------------------

    def _extract_text(self, file_path: Path) -> str:
        """Extract flat text from any supported file format."""
        file_ext = file_path.suffix.lower()
        if file_ext == ".pdf":
            return self._extract_pdf_text(file_path)
        elif file_ext in (".docx", ".doc"):
            return self._extract_docx_text(file_path)
        elif file_ext in (".txt", ".md", ".rtf"):
            return self._extract_plain_text(file_path)
        elif file_ext in (".xlsx", ".xls"):
            return self._extract_excel_text_fallback(file_path)
        else:
            try:
                return self._extract_plain_text(file_path)
            except Exception:
                raise ValueError(f"Unsupported file type: {file_ext}")

    def _extract_pdf_text(self, file_path: Path) -> str:
        text_parts = []
        try:
            with fitz.open(str(file_path)) as pdf:
                for page_num, page in enumerate(pdf, 1):
                    text = page.get_text()
                    if text.strip():
                        text_parts.append(f"[Page {page_num}]\n{text}")
        except Exception as exc:
            raise ValueError(f"Error parsing PDF: {exc}") from exc
        return "\n\n".join(text_parts)

    def _extract_docx_text(self, file_path: Path) -> str:
        try:
            doc = DocxDocument(str(file_path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = [c.text for c in row.cells if c.text.strip()]
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            return "\n\n".join(paragraphs)
        except Exception as exc:
            raise ValueError(f"Error parsing Word document: {exc}") from exc

    def _extract_plain_text(self, file_path: Path) -> str:
        for encoding in ("utf-8", "latin-1", "cp1252", "ascii"):
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    def _extract_excel_text_fallback(self, file_path: Path) -> str:
        """Best-effort flat-text extraction for Excel when structured path is off."""
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
            lines: List[str] = []
            for name in wb.sheetnames:
                ws = wb[name]
                lines.append(f"[Sheet: {name}]")
                for row in ws.iter_rows(values_only=True):
                    vals = [str(v) for v in row if v is not None]
                    if vals:
                        lines.append(" | ".join(vals))
            wb.close()
            return "\n".join(lines)
        except ImportError:
            raise ValueError(
                "openpyxl is required to read Excel files. "
                "Install with: pip install openpyxl"
            )
        except Exception as exc:
            raise ValueError(f"Error reading Excel file: {exc}") from exc

    # ------------------------------------------------------------------
    # Character-window chunker (original, unchanged)
    # ------------------------------------------------------------------

    def _create_chunks(self, text: str) -> List[Dict[str, Any]]:
        """Create overlapping character-window chunks (original algorithm)."""
        if not text:
            return []

        text_length = len(text)
        chunks: List[Dict[str, Any]] = []
        start = 0
        chunk_id = 0
        last_progress = 0

        while start < text_length:
            progress = int((start / text_length) * 100)
            if progress >= last_progress + 10 or chunk_id % 100 == 0:
                logger.info(
                    "    -> Chunking progress: %d%% (%d chunks)", progress, chunk_id
                )
                last_progress = progress

            end = min(start + self.chunk_size, text_length)

            if end < text_length:
                min_end = start + max(self.chunk_size // 2, 1)
                para_break = text.rfind("\n\n", min_end, end)
                if para_break != -1:
                    end = para_break
                else:
                    for break_char in (". ", "! ", "? ", "\n"):
                        bp = text.rfind(break_char, min_end, end)
                        if bp != -1:
                            end = bp + len(break_char)
                            break

            chunk_text = text[start:end].strip()
            if chunk_text:
                token_count = len(self.tokenizer.encode(chunk_text))
                chunks.append({
                    "chunk_id": chunk_id,
                    "text": chunk_text,
                    "start_pos": start,
                    "end_pos": end,
                    "char_count": len(chunk_text),
                    "token_count": token_count,
                    # Structured fields — empty for plain-text chunks
                    "element_type": "paragraph",
                    "section_path": "",
                    "page_or_sheet": "",
                    "embedding_text": "",
                    "html_or_markdown": "",
                })
                chunk_id += 1

            new_start = max(end - self.chunk_overlap, start + 1)
            if new_start <= start:
                new_start = start + max(1, self.chunk_size // 2)
            start = new_start

            if chunk_id > 10_000:
                logger.warning("    -> Chunk limit reached (%d), stopping", chunk_id)
                break

        logger.info("    -> Created %d chunks", len(chunks))
        return chunks

    # ------------------------------------------------------------------
    # Hashing
    # ------------------------------------------------------------------

    def compute_file_hash(self, file_path: Path) -> str:
        """Calculate SHA-256 hash of a file (public)."""
        return self._calculate_file_hash(file_path)

    def _calculate_file_hash(self, file_path: Path) -> str:
        sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for block in iter(lambda: f.read(4096), b""):
                sha256.update(block)
        return sha256.hexdigest()


# ------------------------------------------------------------------
# Parser routing helpers  (module-level)
# ------------------------------------------------------------------

# File types handled by the structured pipeline
_STRUCTURED_TYPES = {".pdf", ".docx", ".doc", ".xlsx", ".xls"}


def _route_to_parser(
    file_path: Path,
    file_type: str,
    visual_config=None,     # VisualConfig | None
) -> List[Dict[str, Any]]:
    """Dispatch to the correct structured parser by file type."""
    if file_type in (".docx", ".doc"):
        from .parsers.structured_docx import parse_docx_structured
        return parse_docx_structured(file_path, visual_config=visual_config)

    if file_type in (".xlsx", ".xls"):
        from .parsers.excel_parser import parse_excel_structured
        # Excel has no visual content — visual_config not used here
        return parse_excel_structured(file_path)

    if file_type == ".pdf":
        from .parsers.pdf_parser import parse_pdf_structured
        return parse_pdf_structured(file_path, visual_config=visual_config)

    raise ValueError(f"No structured parser for file type '{file_type}'")


# ------------------------------------------------------------------
# Utility functions (unchanged public API)
# ------------------------------------------------------------------

def is_supported_file(file_path: Path, extensions: List[str]) -> bool:
    """Check if a file's extension is in the supported list."""
    return file_path.suffix.lower() in extensions


def get_file_info(file_path: Path) -> Optional[Dict[str, Any]]:
    """Return basic file metadata without parsing content."""
    if not file_path.exists():
        return None
    stats = file_path.stat()
    mime_type, _ = mimetypes.guess_type(str(file_path))
    return {
        "path": str(file_path),
        "name": file_path.name,
        "size": stats.st_size,
        "size_mb": round(stats.st_size / (1024 * 1024), 2),
        "modified": datetime.fromtimestamp(stats.st_mtime).isoformat(),
        "created": datetime.fromtimestamp(stats.st_ctime).isoformat(),
        "extension": file_path.suffix.lower(),
        "mime_type": mime_type,
    }
